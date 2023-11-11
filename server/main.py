from docx import Document
from pymongo import MongoClient
from os import listdir, path
import re
from flask import Flask, request, jsonify, abort, Response
from bson import ObjectId
from werkzeug.security import generate_password_hash, check_password_hash
import uuid
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
from bson import json_util

app = Flask(__name__)
CORS(app)
SECRET_KEY = "your_secret_key"
client = MongoClient("mongodb://localhost:27017/")
db = client["Proscom"]
collection = db["TasksData"]
users_collection = db["UsersData"]
articles_collection = db["Articles"]
tests_collection = db["TestsData"]

doc = Document('data/articles.docx')



def insert_roles_and_materials_to_db(roles_and_materials):
    for role, materials_str in roles_and_materials:
        materials_list = [mat.strip() for mat in materials_str.split(',')]
        articles_collection.insert_one({
            "role": role.strip(),
            "articles": materials_list
        })



def read_docx(file_path): client = MongoClient("mongodb://localhost:27017/")


db = client["Proscom"]
collection = db["TasksData"]
users_collection = db["UsersData"]
articles_collection = db["Articles"]

doc = Document('data/articles.docx')


def parse_roles_and_materials(text):
    roles_and_materials = re.findall(r'Роль:\s*(.+?)\nМатериалы:\s*(.+?)(?=\n\Роль:|$)', text, re.DOTALL)
    return roles_and_materials


def insert_roles_and_materials_to_db(roles_and_materials):
    for role, materials_str in roles_and_materials:
        materials_list = []
        for mat in materials_str.split(','):
            mat = mat.strip()
            random_id = str(uuid.uuid4())
            materials_list.append({"title": mat, "id": random_id})

        articles_collection.insert_one({
            "role": role.strip(),
            "materials": materials_list
        })


full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
roles_and_materials = parse_roles_and_materials(full_text)
articles_collection.delete_many({})
insert_roles_and_materials_to_db(roles_and_materials)


def read_docx(file_path):
    document = Document(file_path)
    return ' '.join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip() != '')


def parse_test_document(text):
    question_pattern = re.compile(
        r'(\d+)\.\s*(.+?)\?'
        r'\s*A\)\s*(.+?)\s*'
        r'B\)\s*(.+?)\s*'
        r'C\)\s*(.+?)\s*'
        r'Правильный ответ:\s*([A-C])',
        re.DOTALL
    )

    matches = question_pattern.findall(text)

    questions = [
        {
            'number': match[0],
            'question': match[1].strip(),
            'options': {
                'A': match[2].strip(),
                'B': match[3].strip(),
                'C': match[4].strip(),
            },
            'answer': match[5],
        }
        for match in matches
    ]

    return questions


mypath = 'data'
onlyfiles = [f for f in listdir(mypath) if path.isfile(path.join(mypath, f))]

for file in onlyfiles:
    if file.endswith('.docx') and not file.endswith('_test.docx'):
        text = read_docx(path.join(mypath, file))
        name = path.splitext(file)[0]
        doc = {
            "name": name,
            "content": text
        }
        test_file = f"{name}_test.docx"
        test_file_path = path.join(mypath, test_file)
        if path.exists(test_file_path):
            test_text = read_docx(test_file_path)
            doc['test'] = parse_test_document(test_text)

        collection.insert_one(doc)

UPLOAD_FOLDER = '/data'
ALLOWED_EXTENSIONS = {'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        abort(400, description="No file part in the request.")
    file = request.files['file']
    if file.filename == '':
        abort(400, description="No selected file.")
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return "File uploaded successfully."


@app.route('/register', methods=['POST'])
def register():
    username = request.json.get('username')
    password = request.json.get('password')
    priority = request.json.get('priority')
    role = request.json.get('role')

    if not username or not password:
        return jsonify({'message': 'Missing username or password'}), 400

    if users_collection.find_one({'username': username}):
        return jsonify({'message': 'Username already exists'}), 409

    password_hash = generate_password_hash(password)

    auth_key = str(uuid.uuid4())

    user_id = users_collection.insert_one({
        'username': username,
        'password': password_hash,
        'priority': priority,
        'role': role,
        'auth_key': auth_key
    }).inserted_id

    return jsonify({'auth_key': auth_key}), 201


@app.route('/login', methods=['POST'])
def login():
    username = request.json.get('username')
    password = request.json.get('password')

    if not username or not password:
        return jsonify({'message': 'Missing username or password'}), 400

    user = users_collection.find_one({'username': username})

    if user and check_password_hash(user['password'], password):
        return jsonify({'auth_key': user['auth_key']}), 200
    else:
        return jsonify({'message': 'Invalid credentials'}), 401


@app.route('/users', methods=['GET'])
def get_users():
    users_cursor = users_collection.find({}, {"username": 1, "role": 1, "priority": 1, "_id": 1})
    users = []
    for user in users_cursor:
        user['_id'] = str(user['_id'])
        users.append(user)
    return Response(json_util.dumps(users), mimetype='application/json')


@app.route('/user', methods=['GET'])
def get_user():
    auth_key = request.args.get('auth_key')

    if not auth_key:
        return jsonify({'message': 'Auth key is missing'}), 401

    user = users_collection.find_one({'auth_key': auth_key})

    if not user:
        return jsonify({'message': 'User not found'}), 404

    user_data = {k: str(v) if isinstance(v, ObjectId) else v for k, v in user.items()}

    del user_data['password']
    del user_data['auth_key']

    return jsonify(user_data)


@app.route('/alltasks', methods=['GET'])
def get_tasks():
    tasks = list(collection.find())
    tasks_json = []
    for task in tasks:
        task['_id'] = str(task['_id'])
        tasks_json.append(task)

    return jsonify(tasks_json), 200


@app.route('/tasks', methods=['GET'])
def get_tasks_for_role():
    role = request.args.get('role')

    if not role:
        return jsonify({"error": "Role is missing"}), 400

    role_record = articles_collection.find_one({"role": role})

    if role_record and 'materials' in role_record:
        ids_and_titles = []
        for material in role_record['materials']:
            task = collection.find_one({"name": material['title']})
            if task:
                material["_id"] = str(task["_id"])
            ids_and_titles.append(material)

        titles = [material['title'] for material in role_record['materials']]
        tasks = list(collection.find({"name": {"$in": titles}}))

        tasks_list = []

        for i, task in enumerate(tasks):
            mapped_id = next((item for item in ids_and_titles if item["title"] == task["name"]), None)

            if mapped_id:
                tasks_list.append({
                    "name": task["name"],
                    "content": task.get("content"),
                    "test": task.get("test"),
                    "_id": mapped_id["_id"]
                })

        return jsonify(tasks_list), 200
    else:
        return jsonify({"message": f"No tasks found for role: {role}"}), 404


@app.route('/tasks/<id>', methods=['GET'])
def get_task_by_id(id):
    try:
        task = collection.find_one({"_id": ObjectId(id)})
        if task:
            task['_id'] = str(task['_id'])
            return jsonify(task), 200
        else:
            return jsonify({"message": "Задача не найдена"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 400


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=8010, debug=True)
